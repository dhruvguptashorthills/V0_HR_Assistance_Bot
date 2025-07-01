import os
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
import copy
import re
import io
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

class DocxUtils:
    @staticmethod
    def add_background_image(doc, bg_image_path="templates/bg.png", opacity=0.17):
        """Add a full-page background image as the first element with specified opacity"""
        try:
            # Verify image file exists
            if not os.path.exists(bg_image_path):
                print(f"Error: Background image file not found at {bg_image_path}")
                # Attempt VML fallback if image path is invalid
                return DocxUtils.add_background_image_vml(doc, bg_image_path, opacity)

            section = doc.sections[0]
            header = section.header

            # Clear all existing header content to ensure image is first
            header._element.clear_content()
            for para in header.paragraphs:
                try:
                    p = para._element
                    p.getparent().remove(p)
                except:
                    para.clear()

            # Create a new paragraph for the background image
            bg_para = header.add_paragraph()
            bg_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            bg_para.paragraph_format.space_before = Pt(0)
            bg_para.paragraph_format.space_after = Pt(0)

            # Get page dimensions (US Letter: 8.5" x 11")
            page_width = section.page_width.inches
            page_height = section.page_height.inches

            try:
                run = bg_para.add_run()
                # Add image, scaling to page width
                picture = run.add_picture(bg_image_path, width=Inches(page_width))

                # Adjust height to cover page, maintaining aspect ratio if possible
                img_width, img_height = picture.width, picture.height
                if img_width > 0:
                    aspect_ratio = img_height / img_width
                    target_height = page_width * aspect_ratio
                    picture.height = Inches(min(target_height, page_height))
                    if target_height > page_height:
                        picture.width = Inches(page_height / aspect_ratio)

                # Access the drawing element
                drawing_elements = run._element.xpath('.//wp:anchor | .//wp:inline')
                if not drawing_elements:
                    print("Error: No drawing element found for image, trying VML fallback")
                    return DocxUtils.add_background_image_vml(doc, bg_image_path, opacity)

                drawing = drawing_elements[0]

                # Ensure image is behind all content
                if drawing.tag.endswith('anchor'):
                    drawing.set('behindDoc', '1')
                    drawing.set('relativeHeight', '0')

                    # Position at top-left of page
                    position_h = drawing.xpath('.//wp:positionH')
                    if position_h:
                        position_h[0].set('relativeFrom', 'page')
                        pos_h_offset = position_h[0].xpath('.//wp:posOffset')
                        if pos_h_offset:
                            pos_h_offset[0].text = '0'

                    position_v = drawing.xpath('.//wp:positionV')
                    if position_v:
                        position_v[0].set('relativeFrom', 'page')
                        pos_v_offset = position_v[0].xpath('.//wp:posOffset')
                        if pos_v_offset:
                            pos_v_offset[0].text = '0'

                # Apply opacity
                blip_elements = drawing.xpath('.//a:blip')
                if blip_elements:
                    blip = blip_elements[0]
                    effect_lst = blip.xpath('.//a:effectLst')
                    if not effect_lst:
                        effect_lst = OxmlElement('a:effectLst')
                        blip.append(effect_lst)
                    else:
                        effect_lst = effect_lst[0]

                    # Remove existing alphaModFix
                    for existing_alpha in effect_lst.xpath('.//a:alphaModFix'):
                        effect_lst.remove(existing_alpha)

                    # Add new alpha modulation
                    alpha_mod_fix = OxmlElement('a:alphaModFix')
                    alpha_mod_fix.set('amt', str(int(opacity * 100000)))
                    effect_lst.append(alpha_mod_fix)
                else:
                    print("Error: No blip element found for opacity, trying VML fallback")
                    return DocxUtils.add_background_image_vml(doc, bg_image_path, opacity)

                print("Background image added successfully via primary method")
                return True

            except Exception as e:
                print(f"Error adding background image via primary method: {e}")
                print("Attempting VML fallback")
                return DocxUtils.add_background_image_vml(doc, bg_image_path, opacity)

        except Exception as e:
            print(f"Critical error in add_background_image: {e}")
            print("Attempting VML fallback as last resort")
            return DocxUtils.add_background_image_vml(doc, bg_image_path, opacity)

    @staticmethod
    def add_background_image_vml(doc, bg_image_path="templates/bg.png", opacity=0.17):
        """Fallback: Add background image using VML for compatibility"""
        try:
            if not os.path.exists(bg_image_path):
                print(f"Error: Background image file not found at {bg_image_path} for VML")
                return False

            section = doc.sections[0]
            sect_pr = section._sectPr

            # Add VML background
            vml_background = OxmlElement('w:background')
            vml_background.set(qn('w:color'), 'FFFFFF')  # White fallback

            vml_shape = OxmlElement('v:background')
            vml_shape.set('id', '_x0000_s1024')

            vml_fill = OxmlElement('v:fill')
            vml_fill.set('type', 'tile')
            vml_fill.set('src', bg_image_path)
            vml_fill.set('opacity', str(opacity))

            vml_shape.append(vml_fill)
            vml_background.append(vml_shape)
            sect_pr.append(vml_background)

            print("Background image added successfully via VML")
            return True
        except Exception as e:
            print(f"Error in add_background_image_vml: {e}")
            return False

    # Include the rest of the DocxUtils class methods here (unchanged)
    # For brevity, only the modified methods are shown above
    # Add the original methods like clean_na_values, generate_docx, etc., as needed

    @staticmethod
    def generate_docx(data, keywords=None, left_logo_path="templates/left_logo SMALL.png", right_logo_path="templates/right_logo_small.png"):
        """
        Generate a .docx resume matching the PDF template exactly.
        Returns a BytesIO object containing the Word file.
        """
        # Create a new document
        doc = Document()

        # Add background image as the very first operation
        DocxUtils.add_background_image(doc, bg_image_path="templates/bg.png", opacity=0.17)

        # Set page margins with small values for compatibility
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.2)
            section.bottom_margin = Inches(0.2)
            section.left_margin = Inches(0.2)
            section.right_margin = Inches(0.2)
            section.header_distance = Inches(0.15)
            section.footer_distance = Inches(0.15)
            section.page_width = Inches(8.5)
            section.page_height = Inches(11)

        # Add robust page border
        DocxUtils.add_robust_page_border(doc)

        # Create a deep copy and clean NA values
        data_copy = copy.deepcopy(data)
        data_copy = DocxUtils.clean_na_values(data_copy)
        
        # Limit skills to prevent left column overflow (max 18 skills for DOCX)
        if data_copy.get('skills') and len(data_copy['skills']) > 18:
            data_copy['skills'] = data_copy['skills'][:18]
        
        # Limit certifications to prevent overflow (max 5 certifications)
        if data_copy.get('certifications') and len(data_copy['certifications']) > 5:
            data_copy['certifications'] = data_copy['certifications'][:5]

        # Header with logos
        header = doc.sections[0].header
        header.is_linked_to_previous = False

        # Create header table after background image
        header_table = DocxUtils.create_compatible_table(header, rows=1, cols=3, width_inches=8.1)
        header_table.columns[0].width = Inches(2.7)
        header_table.columns[1].width = Inches(2.7)
        header_table.columns[2].width = Inches(2.7)

        # Left logo
        left_cell = header_table.cell(0, 0)
        left_para = left_cell.paragraphs[0]
        left_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        DocxUtils.add_word_optimized_spacing(left_para, space_before=2, space_after=0)
        left_para.paragraph_format.left_indent = Pt(12)
        try:
            left_run = left_para.add_run()
            left_run.add_picture(left_logo_path, height=Inches(0.35))
        except Exception:
            left_run = left_para.add_run("ShorthillsAI")
            DocxUtils.add_word_font_optimization(left_run, 'Montserrat', 10, True, RGBColor(242, 93, 93))

        # Right logo
        right_cell = header_table.cell(0, 2)
        right_para = right_cell.paragraphs[0]
        right_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        DocxUtils.add_word_optimized_spacing(right_para, space_before=2, space_after=0)
        right_para.paragraph_format.right_indent = Pt(12)
        try:
            right_run = right_para.add_run()
            right_run.add_picture(right_logo_path, height=Inches(0.45))
        except Exception:
            right_run = right_para.add_run("Microsoft Partner")
            right_run.font.name = 'Montserrat'
            right_run.font.size = Pt(10)
            right_run.font.color.rgb = RGBColor(102, 102, 102)

        # Main content table (rest of the method remains unchanged)
        main_table = DocxUtils.create_compatible_table(doc, rows=1, cols=2, width_inches=8.1)
        DocxUtils.set_fixed_column_widths(main_table, 2.8, 5.3)
        DocxUtils.ensure_table_column_borders(main_table, 0, 'CCCCCC')

        left_cell = main_table.cell(0, 0)
        right_cell = main_table.cell(0, 1)
        left_cell._tc.clear_content()
        right_cell._tc.clear_content()

        # Left column padding
        left_cell_padding = left_cell.add_paragraph()
        left_cell_padding.paragraph_format.left_indent = Pt(12)

        # Name and Title
        name_title_table = DocxUtils.create_compatible_table(left_cell, rows=2, cols=1, width_inches=2.8)
        name_title_table.columns[0].width = Inches(2.8)

        name_cell = name_title_table.cell(0, 0)
        DocxUtils.add_grey_background(name_cell)
        name_para = name_cell.paragraphs[0]
        DocxUtils.add_word_optimized_spacing(name_para, space_before=0, space_after=2)
        name_para.paragraph_format.left_indent = Pt(8)
        name_run = name_para.add_run(data_copy.get('name', ''))
        DocxUtils.add_word_font_optimization(name_run, 'Montserrat', 17, True, RGBColor(242, 93, 93))

        title_cell = name_title_table.cell(1, 0)
        DocxUtils.add_grey_background(title_cell)
        title_para = title_cell.paragraphs[0]
        DocxUtils.add_word_optimized_spacing(title_para, space_after=10)
        title_para.paragraph_format.left_indent = Pt(8)
        title_parts = DocxUtils.clean_html_text(data_copy.get('title', ''))
        for text, is_bold in title_parts:
            if text.strip():
                title_run = title_para.add_run(text)
                DocxUtils.add_word_font_optimization(title_run, 'Montserrat', 13, True, RGBColor(34, 34, 34))

        # Skills Section
        if data_copy.get('skills'):
            skills_para = left_cell.add_paragraph()
            DocxUtils.add_word_optimized_spacing(skills_para, space_before=10, space_after=2)
            skills_para.paragraph_format.left_indent = Pt(12)
            skills_run = skills_para.add_run('SKILLS')
            DocxUtils.add_word_font_optimization(skills_run, 'Montserrat', 11, True, RGBColor(242, 93, 93))
            
            for skill in data_copy['skills']:
                skill_parts = DocxUtils.clean_html_text(skill)
                skill_para = left_cell.add_paragraph()
                DocxUtils.add_word_optimized_spacing(skill_para, space_after=3)
                skill_para.paragraph_format.left_indent = Pt(28)
                arrow_run = skill_para.add_run("▶ ")
                DocxUtils.add_word_font_optimization(arrow_run, 'Montserrat', 11, True, RGBColor(242, 93, 93))
                for text, is_bold in skill_parts:
                    if text.strip():
                        skill_run = skill_para.add_run(text)
                        DocxUtils.add_word_font_optimization(skill_run, 'Montserrat', 10, is_bold, RGBColor(34, 34, 34))

        # Education Section
        if data_copy.get('education'):
            edu_para = left_cell.add_paragraph()
            DocxUtils.add_word_optimized_spacing(edu_para, space_before=10, space_after=2)
            edu_para.paragraph_format.left_indent = Pt(12)
            edu_run = edu_para.add_run('EDUCATION')
            DocxUtils.add_word_font_optimization(edu_run, 'Montserrat', 11, True, RGBColor(242, 93, 93))
            
            for edu in data_copy['education']:
                if isinstance(edu, dict):
                    para = left_cell.add_paragraph()
                    DocxUtils.add_word_optimized_spacing(para, space_after=3)
                    para.paragraph_format.left_indent = Pt(28)
                    arrow_run = para.add_run("▶ ")
                    DocxUtils.add_word_font_optimization(arrow_run, 'Montserrat', 11, True, RGBColor(242, 93, 93))
                    if edu.get('degree'):
                        degree_parts = DocxUtils.clean_html_text(edu['degree'])
                        for text, is_bold in degree_parts:
                            if text.strip():
                                degree_run = para.add_run(text)
                                DocxUtils.add_word_font_optimization(degree_run, 'Montserrat', 10, is_bold, RGBColor(34, 34, 34))
                    if edu.get('institution'):
                        para.add_run('\n')
                        inst_parts = DocxUtils.clean_html_text(edu['institution'])
                        for text, is_bold in inst_parts:
                            if text.strip():
                                inst_run = para.add_run(text)
                                DocxUtils.add_word_font_optimization(inst_run, 'Montserrat', 10, True, RGBColor(34, 34, 34))

        # Certifications Section
        if data_copy.get('certifications'):
            cert_para = left_cell.add_paragraph()
            DocxUtils.add_word_optimized_spacing(cert_para, space_before=10, space_after=2)
            cert_para.paragraph_format.left_indent = Pt(12)
            cert_run = cert_para.add_run('CERTIFICATIONS')
            DocxUtils.add_word_font_optimization(cert_run, 'Montserrat', 11, True, RGBColor(242, 93, 93))
            
            for cert in data_copy['certifications']:
                para = left_cell.add_paragraph()
                DocxUtils.add_word_optimized_spacing(para, space_after=3)
                para.paragraph_format.left_indent = Pt(28)
                arrow_run = para.add_run("▶ ")
                DocxUtils.add_word_font_optimization(arrow_run, 'Montserrat', 11, True, RGBColor(242, 93, 93))
                if isinstance(cert, dict):
                    has_content = False
                    if cert.get('title'):
                        title_parts = DocxUtils.clean_html_text(cert['title'])
                        for text, is_bold in title_parts:
                            if text.strip():
                                title_run = para.add_run(text)
                                DocxUtils.add_word_font_optimization(title_run, 'Montserrat', 10, is_bold, RGBColor(34, 34, 34))
                                has_content = True
                    if cert.get('issuer') and has_content:
                        para.add_run('\n')
                        issuer_parts = DocxUtils.clean_html_text(cert['issuer'])
                        for text, is_bold in issuer_parts:
                            if text.strip():
                                issuer_run = para.add_run(text)
                                DocxUtils.add_word_font_optimization(issuer_run, 'Montserrat', 10, True, RGBColor(34, 34, 34))
                    elif cert.get('issuer') and not has_content:
                        issuer_parts = DocxUtils.clean_html_text(cert['issuer'])
                        for text, is_bold in issuer_parts:
                            if text.strip():
                                issuer_run = para.add_run(text)
                                DocxUtils.add_word_font_optimization(issuer_run, 'Montserrat', 10, True, RGBColor(34, 34, 34))
                                has_content = True
                    if cert.get('year') and has_content:
                        year_run = para.add_run(f"\n{cert['year']}")
                        DocxUtils.add_word_font_optimization(year_run, 'Montserrat', 10, False, RGBColor(34, 34, 34))

        # Right column
        right_cell_padding = right_cell.add_paragraph()
        right_cell_padding.paragraph_format.left_indent = Pt(12)

        # Summary Section
        if data_copy.get('summary'):
            summary_title_para = right_cell.add_paragraph()
            DocxUtils.add_word_optimized_spacing(summary_title_para, space_before=0, space_after=2)
            summary_title_para.paragraph_format.left_indent = Pt(12)
            summary_title_run = summary_title_para.add_run('SUMMARY')
            DocxUtils.add_word_font_optimization(summary_title_run, 'Montserrat', 11, True, RGBColor(242, 93, 93))
            
            summary_para = right_cell.add_paragraph()
            DocxUtils.add_word_optimized_spacing(summary_para, space_after=5)
            summary_para.paragraph_format.left_indent = Pt(12)
            summary_parts = DocxUtils.clean_html_text(data_copy['summary'])
            for text, is_bold in summary_parts:
                if text.strip():
                    summary_run = summary_para.add_run(text)
                    DocxUtils.add_word_font_optimization(summary_run, 'Montserrat', 10, is_bold, RGBColor(34, 34, 34))

        # Projects Section
        if data_copy.get('projects'):
            spacing_para = right_cell.add_paragraph()
            DocxUtils.add_word_optimized_spacing(spacing_para, space_after=3)
            
            section_para = right_cell.add_paragraph()
            DocxUtils.add_word_optimized_spacing(section_para, space_after=2)
            section_para.paragraph_format.left_indent = Pt(12)
            section_run = section_para.add_run('KEY RESPONSIBILITIES:')
            DocxUtils.add_word_font_optimization(section_run, 'Montserrat', 11, True, RGBColor(242, 93, 93))
            
            spacing_para2 = right_cell.add_paragraph()
            DocxUtils.add_word_optimized_spacing(spacing_para2, space_after=1)
            
            for idx, project in enumerate(data_copy['projects']):
                if project.get('title'):
                    proj_title_para = right_cell.add_paragraph()
                    DocxUtils.add_word_optimized_spacing(proj_title_para, space_before=6, space_after=2)
                    proj_title_para.paragraph_format.left_indent = Pt(12)
                    title_run = proj_title_para.add_run(f"Project {idx + 1}: ")
                    DocxUtils.add_word_font_optimization(title_run, 'Montserrat', 11, True, RGBColor(242, 93, 93))
                    title_parts = DocxUtils.clean_html_text(project['title'])
                    for text, is_bold in title_parts:
                        if text.strip():
                            proj_run = proj_title_para.add_run(text)
                            DocxUtils.add_word_font_optimization(proj_run, 'Montserrat', 11, True, RGBColor(242, 93, 93))
                
                if project.get('description'):
                    desc_text = project['description']
                    if isinstance(desc_text, str):
                        desc_text = desc_text.replace('</li>', '</li>\n').replace('•', '\n')
                        bullets = re.split(r'(?<=[.?!])\s+|\n|<br\s*/?>|<li>', desc_text)
                        
                        for bullet_html in bullets:
                            bullet_html = bullet_html.strip()
                            if not bullet_html:
                                continue
                            
                            bullet_parts = DocxUtils.clean_html_text(bullet_html)
                            if not any(part[0].strip() for part in bullet_parts):
                                continue

                            bullet_para = right_cell.add_paragraph()
                            DocxUtils.add_word_optimized_spacing(bullet_para, space_after=3)
                            bullet_para.paragraph_format.left_indent = Pt(28)
                            arrow_run = bullet_para.add_run("▶ ")
                            DocxUtils.add_word_font_optimization(arrow_run, 'Montserrat', 11, True, RGBColor(242, 93, 93))
                            for text, is_bold in bullet_parts:
                                if text.strip():
                                    bullet_run = bullet_para.add_run(text)
                                    DocxUtils.add_word_font_optimization(bullet_run, 'Montserrat', 10, is_bold, RGBColor(34, 34, 34))

        # Footer
        footer = doc.sections[0].footer
        footer.is_linked_to_previous = False
        for para in footer.paragraphs:
            try:
                p = para._element
                p.getparent().remove(p)
            except:
                para.clear()
            
        footer_table = DocxUtils.create_compatible_footer_table(footer, 8.0)
        if footer_table is not None:
            footer_cell = footer_table.cell(0, 0)
            DocxUtils.add_cell_background_compatible(footer_cell, 'F25D5D')
            footer_para = footer_cell.paragraphs[0]
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            DocxUtils.set_standard_spacing(footer_para, space_before_pt=6, space_after_pt=6)
            footer_run = footer_para.add_run("© www.shorthills.ai")
            DocxUtils.apply_standard_font(footer_run, 'Montserrat', 10, False, RGBColor(255, 255, 255))

        # Apply PDF optimizations
        DocxUtils.optimize_for_pdf_export(doc)
        DocxUtils.lock_all_table_layouts(doc)

        # Save document
        docx_file = io.BytesIO()
        doc.save(docx_file)
        docx_file.seek(0)
        return docx_file

    # Include other methods from the original DocxUtils class (unchanged)
    @staticmethod
    def clean_na_values(data):
        """Recursively clean 'NA', 'N/A', empty strings, and None values from resume data."""
        if isinstance(data, dict):
            cleaned = {}
            for key, value in data.items():
                cleaned_value = DocxUtils.clean_na_values(value)
                if cleaned_value is not None and cleaned_value != '':
                    cleaned[key] = cleaned_value
            return cleaned
        elif isinstance(data, list):
            cleaned_list = []
            for item in data:
                cleaned_item = DocxUtils.clean_na_values(item)
                if cleaned_item is not None and cleaned_item != '':
                    if isinstance(cleaned_item, dict) and cleaned_item:
                        cleaned_list.append(cleaned_item)
                    elif not isinstance(cleaned_item, dict):
                        cleaned_list.append(cleaned_item)
            return cleaned_list
        elif isinstance(data, str):
            cleaned_str = data.strip()
            na_values = {'na', 'n/a', 'not applicable', 'not available', 'none', 'null', '-', ''}
            if cleaned_str.lower() in na_values:
                return None
            return cleaned_str
        else:
            return data

    @staticmethod
    def clean_html_text(text):
        """Clean HTML tags and return text with bold formatting info"""
        if not text:
            return []
        
        text = str(text)
        bold_parts = []
        strong_pattern = r'<strong>(.*?)</strong>'
        matches = list(re.finditer(strong_pattern, text, re.IGNORECASE))
        
        if matches:
            current_pos = 0
            for match in matches:
                if match.start() > current_pos:
                    bold_parts.append((text[current_pos:match.start()], False))
                bold_parts.append((match.group(1), True))
                current_pos = match.end()
            if current_pos < len(text):
                bold_parts.append((text[current_pos:], False))
        else:
            bold_parts = [(text, False)]
        
        cleaned_parts = []
        for part_text, is_bold in bold_parts:
            clean_text = re.sub(r'<[^>]+>', '', part_text)
            if clean_text.strip():
                cleaned_parts.append((clean_text, is_bold))
        
        return cleaned_parts

    @staticmethod
    def add_formatted_text(paragraph, text_parts, font_size=10, font_color=RGBColor(34, 34, 34)):
        """Add text with mixed formatting to a paragraph"""
        for text, is_bold in text_parts:
            if text.strip():
                run = paragraph.add_run(text)
                run.font.name = 'Montserrat'
                run.font.size = Pt(font_size)
                run.font.color.rgb = font_color
                if is_bold:
                    run.bold = True

    @staticmethod
    def add_robust_page_border(doc):
        """Add page border with enhanced compatibility"""
        try:
            for section in doc.sections:
                sectPr = section._sectPr
                pgBorders = OxmlElement('w:pgBorders')
                pgBorders.set(qn('w:offsetFrom'), 'page')
                pgBorders.set(qn('w:display'), 'allPages')
                
                for border_name in ['top', 'left', 'bottom', 'right']:
                    border = OxmlElement(f'w:{border_name}')
                    border.set(qn('w:val'), 'single')
                    border.set(qn('w:sz'), '6')
                    border.set(qn('w:space'), '0')
                    border.set(qn('w:color'), 'F25D5D')
                    pgBorders.append(border)
                
                sectPr.append(pgBorders)
        except Exception as e:
            print(f"Could not add page border: {e}")

    @staticmethod
    def create_compatible_table(parent, rows, cols, width_inches=None):
        """Create a table with enhanced compatibility"""
        try:
            if hasattr(parent, '_sectPr') or 'header' in str(type(parent)).lower() or 'footer' in str(type(parent)).lower():
                if width_inches is None:
                    width_inches = 8.0
                table = parent.add_table(rows=rows, cols=cols, width=Inches(width_inches))
            else:
                table = parent.add_table(rows=rows, cols=cols)
        except Exception as e:
            if width_inches is None:
                width_inches = 8.0
            table = parent.add_table(rows=rows, cols=cols, width=Inches(width_inches))
        
        try:
            table.autofit = False
            table.allow_autofit = False
            tbl = table._tbl
            tblPr = tbl.tblPr
            tblLayout = OxmlElement('w:tblLayout')
            tblLayout.set(qn('w:type'), 'fixed')
            tblPr.append(tblLayout)
        except:
            pass
        return table

    @staticmethod
    def add_column_border(cell, border_side='right', color='D3D3D3', width='6'):
        """Add a border to a specific side of a table cell"""
        try:
            tcPr = cell._tc.get_or_add_tcPr()
            tcBorders = tcPr.find(qn('w:tcBorders'))
            if tcBorders is None:
                tcBorders = OxmlElement('w:tcBorders')
                tcPr.append(tcBorders)
            
            border = OxmlElement(f'w:{border_side}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), width)
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), color)
            tcBorders.append(border)
        except Exception as e:
            print(f"Could not add cell border: {e}")

    @staticmethod
    def remove_all_table_borders(table):
        """Remove all table borders"""
        try:
            tbl = table._tbl
            tblPr = tbl.tblPr
            tblBorders = OxmlElement('w:tblBorders')
            for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'nil')
                tblBorders.append(border)
            tblPr.append(tblBorders)
            
            for row in table.rows:
                for cell in row.cells:
                    try:
                        tc = cell._tc
                        tcPr = tc.get_or_add_tcPr()
                        tcBorders = OxmlElement('w:tcBorders')
                        for border_name in ['top', 'left', 'bottom', 'right']:
                            border = OxmlElement(f'w:{border_name}')
                            border.set(qn('w:val'), 'nil')
                            tcBorders.append(border)
                        tcPr.append(tcBorders)
                    except:
                        continue
        except Exception as e:
            print(f"Could not remove table borders: {e}")

    @staticmethod
    def set_standard_spacing(paragraph, space_before_pt=0, space_after_pt=0):
        """Set paragraph spacing"""
        if space_before_pt > 0:
            paragraph.paragraph_format.space_before = Pt(space_before_pt)
        if space_after_pt > 0:
            paragraph.paragraph_format.space_after = Pt(space_after_pt)

    @staticmethod
    def apply_standard_font(run, font_name='Montserrat', font_size_pt=10, is_bold=False, color_rgb=None):
        """Apply font formatting"""
        run.font.name = font_name
        run.font.size = Pt(font_size_pt)
        if is_bold:
            run.bold = True
        if color_rgb:
            run.font.color.rgb = color_rgb

    @staticmethod
    def add_cell_background_compatible(cell, color_hex='F2F2F2'):
        """Add cell background with compatibility"""
        try:
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), color_hex)
            cell._tc.get_or_add_tcPr().append(shading_elm)
            
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_mar = OxmlElement('w:tcMar')
            for margin in ['top', 'left', 'bottom', 'right']:
                mar_elem = OxmlElement(f'w:{margin}')
                mar_elem.set(qn('w:w'), '144')
                mar_elem.set(qn('w:type'), 'dxa')
                tc_mar.append(mar_elem)
            tc_pr.append(tc_mar)
            return True
        except Exception as e:
            print(f"Could not add cell background: {e}")
            return False

    @staticmethod
    def create_compatible_footer_table(footer, width_inches=8.5):
        """Create a footer table"""
        try:
            footer_table = footer.add_table(rows=1, cols=1)
            footer_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            footer_table.autofit = False
            footer_table.columns[0].width = Inches(width_inches)
            DocxUtils.remove_all_table_borders(footer_table)
            return footer_table
        except Exception as e:
            print(f"Could not create footer table: {e}")
            return None

    @staticmethod
    def ensure_table_column_borders(table, column_index=0, border_color='CCCCCC'):
        """Ensure table has proper column borders"""
        try:
            if column_index < len(table.columns) - 1:
                for row in table.rows:
                    cell = row.cells[column_index]
                    DocxUtils.add_column_border(cell, 'right', border_color, '8')
        except Exception as e:
            print(f"Could not add column borders: {e}")

    @staticmethod
    def set_fixed_column_widths(table, left_width_inches, right_width_inches):
        """Set fixed column widths"""
        try:
            table.columns[0].width = Inches(left_width_inches)
            table.columns[1].width = Inches(right_width_inches)
            tbl = table._tbl
            tblPr = tbl.tblPr
            existing_layout = tblPr.find(qn('w:tblLayout'))
            if existing_layout is not None:
                tblPr.remove(existing_layout)
            tblLayout = OxmlElement('w:tblLayout')
            tblLayout.set(qn('w:type'), 'fixed')
            tblPr.append(tblLayout)
            tblW = OxmlElement('w:tblW')
            tblW.set(qn('w:w'), str(int((left_width_inches + right_width_inches) * 1440)))
            tblW.set(qn('w:type'), 'dxa')
            tblPr.append(tblW)
        except Exception as e:
            print(f"Could not set fixed column widths: {e}")

    @staticmethod
    def lock_all_table_layouts(doc):
        """Lock all table layouts to fixed"""
        try:
            for table in doc.tables:
                try:
                    tbl = table._tbl
                    tblPr = tbl.tblPr
                    existing_layout = tblPr.find(qn('w:tblLayout'))
                    if existing_layout is not None:
                        tblPr.remove(existing_layout)
                    tblLayout = OxmlElement('w:tblLayout')
                    tblLayout.set(qn('w:type'), 'fixed')
                    tblPr.append(tblLayout)
                except:
                    continue
            for section in doc.sections:
                try:
                    if section.header:
                        for table in section.header.tables:
                            try:
                                tbl = table._tbl
                                tblPr = tbl.tblPr
                                existing_layout = tblPr.find(qn('w:tblLayout'))
                                if existing_layout is not None:
                                    tblPr.remove(existing_layout)
                                tblLayout = OxmlElement('w:tblLayout')
                                tblLayout.set(qn('w:type'), 'fixed')
                                tblPr.append(tblLayout)
                            except:
                                continue
                    if section.footer:
                        for table in section.footer.tables:
                            try:
                                tbl = table._tbl
                                tblPr = tbl.tblPr
                                existing_layout = tblPr.find(qn('w:tblLayout'))
                                if existing_layout is not None:
                                    tblPr.remove(existing_layout)
                                tblLayout = OxmlElement('w:tblLayout')
                                tblLayout.set(qn('w:type'), 'fixed')
                                tblPr.append(tblLayout)
                            except:
                                continue
                except:
                    continue
        except Exception as e:
            print(f"Could not lock table layouts: {e}")

    @staticmethod
    def optimize_for_pdf_export(doc):
        """Apply optimizations for PDF export"""
        try:
            for section in doc.sections:
                section.different_first_page_header_footer = False
                section.start_type = WD_SECTION.NEW_PAGE
                section.page_width = Inches(8.5)
                section.page_height = Inches(11)
        except Exception as e:
            print(f"Could not optimize for PDF export: {e}")

    @staticmethod
    def add_grey_background(cell):
        """Add grey background to cell"""
        return DocxUtils.add_cell_background_compatible(cell, 'F2F2F2')

    @staticmethod
    def remove_table_borders(table):
        """Remove table borders"""
        return DocxUtils.remove_all_table_borders(table)

    @staticmethod
    def add_page_border(doc):
        """Add page border"""
        return DocxUtils.add_robust_page_border(doc)

    @staticmethod
    def optimize_table_for_word(table):
        """Optimize table for Word"""
        return DocxUtils.remove_all_table_borders(table)

    @staticmethod
    def add_word_optimized_spacing(paragraph, space_before=0, space_after=0, line_spacing=1.0):
        """Add spacing"""
        return DocxUtils.set_standard_spacing(paragraph, space_before, space_after)

    @staticmethod
    def add_word_font_optimization(run, font_name='Montserrat', font_size=10, is_bold=False, color_rgb=None):
        """Font optimization"""
        return DocxUtils.apply_standard_font(run, font_name, font_size, is_bold, color_rgb)